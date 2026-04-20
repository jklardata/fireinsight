from datetime import datetime
from fastapi import FastAPI, Request, Form, UploadFile, File
from fastapi.responses import HTMLResponse, PlainTextResponse, JSONResponse, RedirectResponse, Response
from fastapi.staticfiles import StaticFiles
from fastapi.templating import Jinja2Templates
from typing import Optional
import httpx, uuid, os, json as _json_mod, io, csv as _csv_mod

# ── User uploaded data (session-cookie + /tmp) ──────────────────────────────
_UD_DIR = "/tmp/5alarmdata_uploads"
os.makedirs(_UD_DIR, exist_ok=True)

def _ud_cookie(request: Request) -> str | None:
    return request.cookies.get("ud_id")

def load_user_data(request: Request):
    """Return (incidents, dept_name) from user-uploaded data, or (None, None)."""
    uid = _ud_cookie(request)
    if not uid:
        return None, None
    path = os.path.join(_UD_DIR, f"{uid}.json")
    if not os.path.exists(path):
        return None, None
    try:
        with open(path) as f:
            d = _json_mod.load(f)
        return d["incidents"], d["dept_name"]
    except Exception:
        return None, None

def has_user_data(request: Request) -> bool:
    uid = _ud_cookie(request)
    if not uid:
        return False
    return os.path.exists(os.path.join(_UD_DIR, f"{uid}.json"))

def user_data_summary(request: Request) -> dict | None:
    uid = _ud_cookie(request)
    if not uid:
        return None
    path = os.path.join(_UD_DIR, f"{uid}.json")
    if not os.path.exists(path):
        return None
    try:
        with open(path) as f:
            d = _json_mod.load(f)
        return {"dept_name": d["dept_name"], "count": len(d["incidents"])}
    except Exception:
        return None

# In-memory NERIS API credential store (per server process)
_neris_creds: dict = {"client_id": None, "client_secret": None, "connected": False, "entity": None}

from config import RESEND_API_KEY, DEMO_EMAIL_TO, SUPABASE_URL, SUPABASE_SERVICE_ROLE_KEY, CLERK_PUBLISHABLE_KEY, CLERK_SECRET_KEY, NERIS_CLIENT_ID, NERIS_CLIENT_SECRET
from analytics import summarize_incidents
from insights.trends import generate_trend_summary
from insights.report import generate_chiefs_report
from insights.grant import generate_grant_narrative
from convert.nfirs_to_neris import convert_nfirs_csv, summarise_conversion
from quality.scorer import score_incidents
from insights.quality import generate_quality_narrative

app = FastAPI()
app.mount("/static", StaticFiles(directory="static"), name="static")
templates = Jinja2Templates(directory="templates")

# ── Auth helpers ─────────────────────────────────────────────────────────────

PROTECTED_PATHS = {"/app", "/generate", "/trends-data", "/quality", "/convert", "/map-data"}

async def verify_clerk_session(session_token: str) -> bool:
    """Verify a Clerk __session JWT via Clerk's backend API."""
    if not CLERK_SECRET_KEY or not session_token:
        return False
    try:
        async with httpx.AsyncClient(timeout=5) as client:
            r = await client.post(
                "https://api.clerk.com/v1/tokens/verify",
                headers={"Authorization": f"Bearer {CLERK_SECRET_KEY}"},
                json={"token": session_token},
            )
            return r.status_code == 200
    except Exception:
        return False

@app.middleware("http")
async def clerk_auth_middleware(request: Request, call_next):
    # Auth disabled — re-enable when Clerk Production mode is set up with real domain
    return await call_next(request)


def has_neris_credentials() -> bool:
    """True when server-level NERIS API credentials are configured."""
    return bool(NERIS_CLIENT_ID and NERIS_CLIENT_SECRET)


def load_incidents(neris_id: str, use_mock: bool, start: str, end: str):
    start_dt = datetime.strptime(start, "%Y-%m-%d") if start else None
    end_dt   = datetime.strptime(end,   "%Y-%m-%d") if end   else None

    if use_mock:
        from mock_data import generate_incidents, DEPT
        incidents = generate_incidents(start=start_dt, end=end_dt)
        dept_name = DEPT["name"]
    else:
        from neris import fetch_incidents, fetch_entity
        incidents = fetch_incidents(neris_id, start=start_dt, end=end_dt)
        try:
            entity    = fetch_entity(neris_id)
            dept_name = entity.get("name", neris_id)
        except Exception:
            dept_name = neris_id

    return incidents, dept_name


def resolve_incidents(request: Request, neris_id: str, use_mock: bool, start: str, end: str):
    """
    Unified data resolution:
      1. User-uploaded session data (highest priority, when not using mock)
      2. Live NERIS API (when credentials configured and not using mock)
      3. Mock/sample data fallback
    """
    user_incidents, user_dept = load_user_data(request)
    if user_incidents and not use_mock:
        return user_incidents, user_dept
    if not use_mock and has_neris_credentials():
        return load_incidents(neris_id, False, start, end)
    return load_incidents(neris_id, True, start, end)


@app.get("/", response_class=HTMLResponse)
async def landing(request: Request):
    return templates.TemplateResponse("landing.html", {"request": request})


def clerk_sign_in_url(publishable_key: str, redirect_url: str = "") -> str:
    """Derive Clerk's hosted sign-in URL from the publishable key."""
    import base64
    from urllib.parse import urlencode
    try:
        b64 = publishable_key.split('_', 2)[2]
        b64 += '=' * ((4 - len(b64) % 4) % 4)
        frontend_api = base64.b64decode(b64).decode().rstrip('$')
        # Account Portal drops ".clerk" from frontend API domain
        # e.g. flexible-tarpon-81.clerk.accounts.dev → flexible-tarpon-81.accounts.dev
        account_portal = frontend_api.replace('.clerk.accounts.dev', '.accounts.dev')
        url = f"https://{account_portal}/sign-in"
        if redirect_url:
            url += '?' + urlencode({'redirect_url': redirect_url})
        return url
    except Exception:
        return "/app"


@app.get("/sign-in", response_class=HTMLResponse)
async def sign_in(request: Request):
    next_path = request.query_params.get("next", "/app")
    return templates.TemplateResponse("sign_in.html", {
        "request": request,
        "clerk_key": CLERK_PUBLISHABLE_KEY,
        "next": next_path,
    })


@app.get("/sign-up", response_class=HTMLResponse)
async def sign_up(request: Request):
    next_path = request.query_params.get("next", "/app")
    return templates.TemplateResponse("sign_up.html", {
        "request": request,
        "clerk_key": CLERK_PUBLISHABLE_KEY,
        "next": next_path,
    })


@app.get("/sign-out", response_class=HTMLResponse)
async def sign_out(request: Request):
    return templates.TemplateResponse("sign_out.html", {
        "request": request,
        "clerk_key": CLERK_PUBLISHABLE_KEY,
    })


@app.get("/app", response_class=HTMLResponse)
async def app_dashboard(request: Request):
    return templates.TemplateResponse("index.html", {
        "request":         request,
        "clerk_key":       CLERK_PUBLISHABLE_KEY,
        "active_tab":      "home",
        "user_data":       user_data_summary(request),
        "neris_api_ready": has_neris_credentials(),
        "home_tab":        request.query_params.get("tab", "upload"),
        "connect_error":   request.query_params.get("connect_error", ""),
    })


@app.get("/debug-neris")
async def debug_neris(request: Request, neris_id: str = "FD05005639"):
    """Temporary debug endpoint — remove after diagnosis."""
    from fastapi.responses import JSONResponse
    out = {
        "neris_id": neris_id,
        "has_credentials": has_neris_credentials(),
        "client_id_prefix": (NERIS_CLIENT_ID or "")[:8] or "NOT SET",
        "base_url": str(__import__("config").NERIS_BASE_URL),
    }
    if not has_neris_credentials():
        return JSONResponse(out)
    try:
        from neris import get_client
        client = get_client()
        out["client_created"] = True
        result = client.list_incidents(neris_id_entity=neris_id, page_size=5)
        out["result_type"] = str(type(result).__name__)
        try:
            out["result_status"] = result.status_code
            out["result_body"] = result.text[:300]
        except Exception:
            pass
        if isinstance(result, dict):
            out["data_count"] = len(result.get("data", []))
            out["keys"] = list(result.keys())
    except Exception as e:
        out["error"] = str(e)
        out["error_type"] = type(e).__name__
    return JSONResponse(out)


@app.post("/connect-neris")
async def connect_neris(
    request:  Request,
    neris_id: str = Form(...),
    start:    str = Form(""),
    end:      str = Form(""),
):
    import asyncio
    from fastapi.concurrency import run_in_threadpool

    if not has_neris_credentials():
        return RedirectResponse(
            "/app?tab=neris&connect_error=NERIS+API+credentials+not+configured+on+server",
            status_code=303,
        )
    try:
        incidents, dept_name = await asyncio.wait_for(
            run_in_threadpool(load_incidents, neris_id.strip(), False, start, end),
            timeout=25.0,
        )
        if not incidents:
            return RedirectResponse(
                f"/app?tab=neris&connect_error=No+incidents+found+for+{neris_id}",
                status_code=303,
            )
        uid = str(uuid.uuid4())
        path = os.path.join(_UD_DIR, f"{uid}.json")
        with open(path, "w") as f:
            _json_mod.dump({"incidents": incidents, "dept_name": dept_name or neris_id}, f)
        response = RedirectResponse("/app?uploaded=1", status_code=303)
        response.set_cookie("ud_id", uid, max_age=86400 * 7, httponly=True, samesite="lax")
        return response
    except asyncio.TimeoutError:
        return RedirectResponse(
            "/app?tab=neris&connect_error=Request+timed+out+after+25s.+Check+your+NERIS+ID+and+credentials.",
            status_code=303,
        )
    except Exception as e:
        err = str(e)[:120].replace("&", "and").replace("#", "")
        return RedirectResponse(f"/app?tab=neris&connect_error={err}", status_code=303)


@app.post("/generate", response_class=HTMLResponse)
async def generate(
    request: Request,
    report_type: str  = Form(...),
    neris_id:    str  = Form("MOCK-001"),
    use_mock:    bool = Form(False),
    start:       str  = Form(""),
    end:         str  = Form(""),
    period:      str  = Form(""),
    grant_type:  str  = Form("AFG"),
    grant_request: str = Form(""),
):
    error = None
    result = None

    try:
        incidents, dept_name = resolve_incidents(request, neris_id, use_mock, start, end)

        if not incidents:
            error = "No incidents found for the given parameters."
        else:
            stats = summarize_incidents(incidents)
            period_label = period or ("the selected period" if start or end else "the past year")

            if report_type == "trends":
                result = generate_trend_summary(dept_name, stats, period_label)
            elif report_type == "report":
                result = generate_chiefs_report(dept_name, stats, period_label)
            elif report_type == "grant":
                if not grant_request:
                    error = "Please describe what you are requesting for the grant narrative."
                else:
                    result = generate_grant_narrative(
                        dept_name=dept_name,
                        stats=stats,
                        grant_type=grant_type,
                        request_description=grant_request,
                        period=period_label,
                    )
    except Exception as e:
        error = str(e)

    return templates.TemplateResponse("index.html", {
        "request": request,
        "result": result,
        "error": error,
        "user_data": user_data_summary(request),
        "form": {
            "report_type": report_type,
            "neris_id": neris_id,
            "use_mock": use_mock,
            "start": start,
            "end": end,
            "period": period,
            "grant_type": grant_type,
            "grant_request": grant_request,
        },
    })


@app.post("/trends-data", response_class=HTMLResponse)
async def trends_data(
    request:  Request,
    neris_id: str  = Form("MOCK-001"),
    use_mock: bool = Form(False),
    start:    str  = Form(""),
    end:      str  = Form(""),
    period:   str  = Form(""),
):
    error  = None
    trends = None
    try:
        incidents, dept_name = resolve_incidents(request, neris_id, use_mock, start, end)
        if not incidents:
            error = "No incidents found for the given parameters."
        else:
            stats = summarize_incidents(incidents)
            period_label = period or ("the selected period" if start or end else "the past year")

            by_month = stats.get("by_month", {})
            months   = sorted(by_month.keys())[-12:]
            month_counts = [by_month[m] for m in months]

            by_year: dict = {}
            for m, c in by_month.items():
                yr = m.split()[-1] if " " in m else m[:4]
                by_year[yr] = by_year.get(yr, 0) + c

            call_types = stats.get("call_types", {})
            rt = stats.get("response_time_seconds", {})
            avg_rt = rt.get("average")
            avg_rt_fmt = (f"{int(avg_rt) // 60}m {int(avg_rt) % 60:02d}s") if avg_rt else None

            trends = {
                "dept_name":    dept_name,
                "period":       period_label,
                "total":        len(incidents),
                "stats":        stats,
                "months":       months,
                "month_counts": month_counts,
                "by_year":      by_year,
                "call_types":   call_types,
                "top_types":    list(call_types.items())[:5],
                "avg_rt":       avg_rt,
                "avg_rt_fmt":   avg_rt_fmt,
                "rt_sample":    rt.get("sample_size", 0),
                "busiest_days": stats.get("busiest_days", []),
            }
    except Exception as e:
        error = str(e)

    return templates.TemplateResponse("index.html", {
        "request":     request,
        "trends_data": trends,
        "error":       error,
        "active_tab":  "trends",
        "user_data":   user_data_summary(request),
    })


@app.post("/upload-data")
async def upload_data(request: Request, incident_file: UploadFile = File(...)):
    try:
        raw = await incident_file.read()
        dept_name = incident_file.filename.rsplit(".", 1)[0].replace("_", " ").title()
        try:
            text = raw.decode("utf-8-sig")
            reader = _csv_mod.DictReader(io.StringIO(text))
            incidents = [row for row in reader]
        except Exception:
            incidents = _json_mod.loads(raw)
        if not incidents:
            return RedirectResponse("/app?upload_error=empty", status_code=303)
        uid = str(uuid.uuid4())
        path = os.path.join(_UD_DIR, f"{uid}.json")
        with open(path, "w") as f:
            _json_mod.dump({"incidents": incidents, "dept_name": dept_name}, f)
        response = RedirectResponse("/app?uploaded=1", status_code=303)
        response.set_cookie("ud_id", uid, max_age=86400 * 7, httponly=True, samesite="lax")
        return response
    except Exception as e:
        return RedirectResponse(f"/app?upload_error={str(e)[:60]}", status_code=303)


@app.get("/clear-data")
async def clear_data(request: Request):
    uid = _ud_cookie(request)
    if uid:
        path = os.path.join(_UD_DIR, f"{uid}.json")
        if os.path.exists(path):
            os.remove(path)
    response = RedirectResponse("/app", status_code=303)
    response.delete_cookie("ud_id")
    return response


@app.get("/start", response_class=HTMLResponse)
async def onboarding(request: Request):
    return templates.TemplateResponse("onboarding.html", {"request": request})


@app.get("/about", response_class=HTMLResponse)
async def about(request: Request):
    return templates.TemplateResponse("about.html", {"request": request})


@app.post("/download")
async def download(content: str = Form(...), filename: str = Form("5alarmdata-report.md")):
    return PlainTextResponse(
        content,
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
        media_type="text/markdown",
    )


@app.post("/download-docx")
async def download_docx(content: str = Form(...), filename: str = Form("5alarmdata-report")):
    from docx import Document
    from docx.shared import Pt, RGBColor
    import re as _re

    doc = Document()
    # Style the default paragraph font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    for line in content.splitlines():
        line = line.rstrip()
        if line.startswith("### "):
            p = doc.add_heading(line[4:], level=3)
        elif line.startswith("## "):
            p = doc.add_heading(line[3:], level=2)
        elif line.startswith("# "):
            p = doc.add_heading(line[2:], level=1)
        elif line.startswith("- ") or line.startswith("* "):
            doc.add_paragraph(line[2:], style="List Bullet")
        elif _re.match(r"^\d+\. ", line):
            doc.add_paragraph(_re.sub(r"^\d+\. ", "", line), style="List Number")
        elif line == "---" or line == "***":
            doc.add_paragraph("─" * 60)
        elif line:
            doc.add_paragraph(line)
        else:
            doc.add_paragraph("")

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    fname = f"{filename}.docx"
    from fastapi.responses import StreamingResponse
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{fname}"'},
    )


# ── NFIRS → NERIS Converter ─────────────────────────────────────────────────

@app.post("/convert", response_class=HTMLResponse)
async def convert_upload(
    request: Request,
    file: UploadFile = File(...),
):
    error = None
    incidents = None
    summary = None
    warnings = []
    filename = file.filename or "upload.csv"

    try:
        raw_bytes = await file.read()
        # Try common encodings
        csv_text = None
        for enc in ("utf-8", "latin-1", "cp1252"):
            try:
                csv_text = raw_bytes.decode(enc)
                break
            except UnicodeDecodeError:
                continue

        if csv_text is None:
            raise ValueError("Could not decode file. Please export your NFIRS data as UTF-8 or Latin-1 CSV.")

        incidents, warnings = convert_nfirs_csv(csv_text)

        if not incidents:
            error = "No incidents found in the uploaded file. Check that it is a valid NFIRS Basic Module CSV export."
        else:
            summary = summarise_conversion(incidents)

    except Exception as e:
        error = str(e)

    return templates.TemplateResponse("index.html", {
        "request": request,
        "active_tab": "convert",
        "convert": {
            "filename": filename,
            "incidents": incidents,
            "summary": summary,
            "warnings": warnings,
            "error": error,
        },
    })


# ── Data Quality Scorer ──────────────────────────────────────────────────────

@app.post("/quality", response_class=HTMLResponse)
async def quality_score(
    request: Request,
    neris_id: str  = Form("MOCK-001"),
    use_mock: bool = Form(False),
    start:    str  = Form(""),
    end:      str  = Form(""),
    narrative: bool = Form(False),
):
    error = None
    report = None
    ai_narrative = None

    try:
        incidents, dept_name = resolve_incidents(request, neris_id, use_mock, start, end)

        if not incidents:
            error = "No incidents found for the given parameters."
        else:
            report = score_incidents(incidents)
            if narrative:
                ai_narrative = generate_quality_narrative(dept_name, report)

    except Exception as e:
        error = str(e)

    return templates.TemplateResponse("index.html", {
        "request": request,
        "active_tab": "quality",
        "user_data": user_data_summary(request),
        "quality": {
            "dept_name": dept_name if not error else "",
            "report": report,
            "narrative": ai_narrative,
            "error": error,
        },
    })


# ── Map Data ────────────────────────────────────────────────────────────────

@app.get("/map-data")
async def map_data(
    request: Request,
    neris_id: str = "MOCK-001",
    use_mock: str = "true",
    start: str = "",
    end: str = "",
):
    try:
        _use_mock_bool = use_mock.lower() == "true"
        incidents, dept_name = resolve_incidents(request, neris_id, _use_mock_bool, start, end)
        features = []
        for inc in incidents:
            lat = inc.get("latitude") or inc.get("lat")
            lon = inc.get("longitude") or inc.get("lon")
            if not lat or not lon:
                continue
            features.append({
                "lat": float(lat),
                "lon": float(lon),
                "type": inc.get("incident_type") or "Other",
                "id":   inc.get("neris_id_incident") or inc.get("incident_id") or "",
                "time": inc.get("call_create") or "",
                "arrival": inc.get("arrival_time") or "",
            })
        return JSONResponse({"dept": dept_name, "incidents": features})
    except Exception as e:
        return JSONResponse({"error": str(e)}, status_code=500)


# ── EV / Lithium Battery Fire Analysis ───────────────────────────────────────

@app.post("/ev", response_class=HTMLResponse)
async def ev_analysis(
    request:  Request,
    neris_id: str  = Form("MOCK-001"),
    use_mock: bool = Form(False),
    start:    str  = Form(""),
    end:      str  = Form(""),
):
    error     = None
    report    = None
    dept_name = ""

    try:
        from ev.detector import analyze_ev_incidents
        incidents, dept_name = resolve_incidents(request, neris_id, use_mock, start, end)
        if not incidents:
            error = "No incidents found for the given parameters."
        else:
            report = analyze_ev_incidents(incidents)
    except Exception as e:
        error = str(e)

    return templates.TemplateResponse("index.html", {
        "request":    request,
        "active_tab": "ev",
        "user_data":  user_data_summary(request),
        "ev": {
            "dept_name": dept_name,
            "report":    report,
            "error":     error,
        },
    })


# ── Community Risk Dashboard ──────────────────────────────────────────────────

@app.post("/risk", response_class=HTMLResponse)
async def risk_dashboard(
    request:  Request,
    neris_id: str  = Form("MOCK-001"),
    use_mock: bool = Form(False),
    start:    str  = Form(""),
    end:      str  = Form(""),
):
    error     = None
    summary   = None
    zones     = None
    dept_name = ""

    try:
        from risk.scorer import compute_risk_zones, compute_risk_summary
        incidents, dept_name = load_incidents(neris_id, use_mock, start, end)
        if not incidents:
            error = "No incidents found for the given parameters."
        else:
            zones   = compute_risk_zones(incidents)
            summary = compute_risk_summary(zones)
    except Exception as e:
        error = str(e)

    return templates.TemplateResponse("index.html", {
        "request":    request,
        "active_tab": "risk",
        "risk": {
            "dept_name": dept_name,
            "summary":   summary,
            "zones":     zones,
            "error":     error,
        },
    })


# ── Peer Benchmark ───────────────────────────────────────────────────────────

@app.post("/benchmark", response_class=HTMLResponse)
async def benchmark(
    request:    Request,
    neris_id:   str  = Form("MOCK-001"),
    use_mock:   bool = Form(False),
    start:      str  = Form(""),
    end:        str  = Form(""),
    state:      str  = Form("VA"),
    dept_type:  str  = Form("combination"),
    population: int  = Form(25000),
    narrative:  bool = Form(False),
):
    error = None
    result = None
    ai_narrative = None
    dept_name = ""

    try:
        from benchmark.engine import run_benchmark
        from insights.benchmark_narrative import generate_benchmark_narrative

        incidents, dept_name = resolve_incidents(request, neris_id, use_mock, start, end)

        if not incidents:
            error = "No incidents found for the given parameters."
        else:
            stats  = summarize_incidents(incidents)
            result = run_benchmark(
                dept_stats=stats,
                dept_name=dept_name,
                state=state,
                dept_type=dept_type,
                population=population,
            )
            if narrative:
                ai_narrative = generate_benchmark_narrative(result)

    except Exception as e:
        error = str(e)

    return templates.TemplateResponse("index.html", {
        "request": request,
        "active_tab": "benchmark",
        "user_data": user_data_summary(request),
        "benchmark": {
            "result":    result,
            "narrative": ai_narrative,
            "error":     error,
            "form": {
                "state":      state,
                "dept_type":  dept_type,
                "population": population,
            },
        },
    })


# ── NERIS Compliance Tracker ─────────────────────────────────────────────────

@app.post("/compliance", response_class=HTMLResponse)
async def compliance_check(
    request:       Request,
    use_mock:      bool       = Form(False),
    incident_file: UploadFile = File(None),
):
    error     = None
    report    = None
    dept_name = ""

    try:
        from compliance.checker import check_compliance
        if not use_mock and incident_file and incident_file.filename:
            raw = await incident_file.read()
            dept_name = incident_file.filename.rsplit(".", 1)[0]
            try:
                text = raw.decode("utf-8-sig")
                reader = _csv_mod.DictReader(io.StringIO(text))
                incidents = [row for row in reader]
            except Exception:
                incidents = _json_mod.loads(raw)
        elif not use_mock:
            incidents, dept_name = resolve_incidents(request, "MOCK-001", False, "", "")
        else:
            incidents, dept_name = load_incidents("MOCK-001", True, "", "")
        if not incidents:
            error = "No incidents found in the uploaded file."
        else:
            report = check_compliance(incidents)
    except Exception as e:
        error = str(e)

    return templates.TemplateResponse("index.html", {
        "request":    request,
        "active_tab": "compliance",
        "user_data":  user_data_summary(request),
        "compliance": {
            "dept_name": dept_name,
            "report":    report,
            "error":     error,
        },
        "clerk_key": CLERK_PUBLISHABLE_KEY,
    })


# ── ISO Evidence Pack ─────────────────────────────────────────────────────────

@app.post("/iso", response_class=HTMLResponse)
async def iso_evidence(
    request:   Request,
    neris_id:  str  = Form("MOCK-001"),
    use_mock:  bool = Form(False),
    start:     str  = Form(""),
    end:       str  = Form(""),
    period:    str  = Form(""),
    narrative: bool = Form(False),
):
    error       = None
    iso_metrics = None
    ai_narrative = None
    dept_name   = ""

    try:
        from insights.iso_narrative import compute_iso_metrics, generate_iso_narrative
        incidents, dept_name = resolve_incidents(request, neris_id, use_mock, start, end)
        if not incidents:
            error = "No incidents found for the given parameters."
        else:
            stats       = summarize_incidents(incidents)
            iso_metrics = compute_iso_metrics(incidents)
            period_label = period or ("the selected period" if start or end else "the past year")
            if narrative and iso_metrics:
                ai_narrative = generate_iso_narrative(dept_name, stats, iso_metrics, period_label)
    except Exception as e:
        error = str(e)

    return templates.TemplateResponse("index.html", {
        "request":    request,
        "active_tab": "iso",
        "user_data":  user_data_summary(request),
        "iso": {
            "dept_name":   dept_name,
            "iso_metrics": iso_metrics,
            "narrative":   ai_narrative,
            "error":       error,
        },
    })


# ── Staffing Justification Report ────────────────────────────────────────────

@app.post("/staffing", response_class=HTMLResponse)
async def staffing_report(
    request:   Request,
    neris_id:  str  = Form("MOCK-001"),
    use_mock:  bool = Form(False),
    start:     str  = Form(""),
    end:       str  = Form(""),
    narrative: bool = Form(False),
):
    error       = None
    report      = None
    ai_narrative = None
    dept_name   = ""

    try:
        from staffing.analyzer import analyze_staffing
        from insights.staffing_narrative import generate_staffing_narrative
        incidents, dept_name = resolve_incidents(request, neris_id, use_mock, start, end)
        if not incidents:
            error = "No incidents found for the given parameters."
        else:
            report = analyze_staffing(incidents)
            if narrative and report:
                ai_narrative = generate_staffing_narrative(dept_name, report)
    except Exception as e:
        error = str(e)

    return templates.TemplateResponse("index.html", {
        "request":    request,
        "active_tab": "staffing",
        "user_data":  user_data_summary(request),
        "staffing": {
            "dept_name": dept_name,
            "report":    report,
            "narrative": ai_narrative,
            "error":     error,
        },
    })


# ── NERIS Submit & Validate ──────────────────────────────────────────────────

@app.post("/neris/pre-validate", response_class=HTMLResponse)
async def neris_pre_validate(
    request:       Request,
    neris_id:      str  = Form("MOCK-001"),
    use_mock:      bool = Form(False),
    dept_neris_id: str  = Form(""),
    start:         str  = Form(""),
    end:           str  = Form(""),
):
    error    = None
    pre_val  = None
    try:
        incidents, dept_name = resolve_incidents(request, neris_id, use_mock, start, end)
        if not incidents:
            error = "No incidents found for the given parameters."
        else:
            from submission.pre_validator import pre_validate
            pre_val = pre_validate(incidents, dept_neris_id or None)
            pre_val["dept_name"] = dept_name
    except Exception as e:
        error = str(e)

    return templates.TemplateResponse("index.html", {
        "request":    request,
        "pre_val":    pre_val,
        "error":      error,
        "active_tab": "submit",
        "user_data":  user_data_summary(request),
    })


@app.post("/neris/api-validate", response_class=JSONResponse)
async def neris_api_validate(
    neris_id:      str = Form(...),
    dept_neris_id: str = Form(...),
    use_mock:      bool = Form(False),
    start:         str = Form(""),
    end:           str = Form(""),
):
    """Call NERIS validate_incident for each incident (requires API credentials)."""
    if not _neris_creds.get("connected"):
        return JSONResponse({"ok": False, "message": "Not connected to NERIS API. Enter your credentials first."}, status_code=401)
    try:
        incidents, _ = load_incidents(neris_id, use_mock, start, end)
        from neris_api_client import NerisApiClient
        from neris_api_client.config import Config, GrantType
        cfg = Config(
            base_url="https://api.neris.fsri.org/v1",
            grant_type=GrantType.CLIENT_CREDENTIALS,
            client_id=_neris_creds["client_id"],
            client_secret=_neris_creds["client_secret"],
        )
        client = NerisApiClient(cfg)
        results = []
        for inc in incidents[:50]:  # cap at 50 for the validate preview
            try:
                body = {
                    "base": {
                        "department_neris_id": dept_neris_id,
                        "incident_number": str(inc.get("neris_id_incident") or inc.get("incident_id") or ""),
                        "location": {"address": inc.get("address", "")},
                    },
                    "incident_types": [{"type": inc.get("incident_type", "OTHER")}],
                    "dispatch": {
                        "incident_number": str(inc.get("neris_id_incident") or inc.get("incident_id") or ""),
                        "call_arrival": inc.get("call_create", ""),
                        "call_answered": inc.get("call_create", ""),
                    },
                }
                resp = client.validate_incident(dept_neris_id, body)
                results.append({"incident_id": inc.get("neris_id_incident"), "ok": True, "detail": resp})
            except Exception as ex:
                results.append({"incident_id": inc.get("neris_id_incident"), "ok": False, "detail": str(ex)})
        ok_count = sum(1 for r in results if r["ok"])
        return JSONResponse({"ok": True, "validated": len(results), "passed": ok_count, "results": results})
    except Exception as e:
        return JSONResponse({"ok": False, "message": str(e)}, status_code=500)


@app.post("/neris/submit", response_class=JSONResponse)
async def neris_submit(
    neris_id:      str  = Form(...),
    dept_neris_id: str  = Form(...),
    use_mock:      bool = Form(False),
    start:         str  = Form(""),
    end:           str  = Form(""),
    dry_run:       bool = Form(False),
):
    """Submit incidents to NERIS API. Requires API credentials."""
    if not _neris_creds.get("connected"):
        return JSONResponse({"ok": False, "message": "Not connected to NERIS API."}, status_code=401)
    try:
        incidents, dept_name = load_incidents(neris_id, use_mock, start, end)
        from submission.pre_validator import pre_validate
        pre = pre_validate(incidents, dept_neris_id)
        submittable = [
            inc for inc, res in zip(incidents, pre["incidents"])
            if res["status"] in ("ready", "warning")
        ]
        if dry_run:
            return JSONResponse({
                "ok": True, "dry_run": True,
                "would_submit": len(submittable),
                "blocked": pre["with_errors"],
                "dept_name": dept_name,
            })
        from neris_api_client import NerisApiClient
        from neris_api_client.config import Config, GrantType
        cfg = Config(
            base_url="https://api.neris.fsri.org/v1",
            grant_type=GrantType.CLIENT_CREDENTIALS,
            client_id=_neris_creds["client_id"],
            client_secret=_neris_creds["client_secret"],
        )
        client = NerisApiClient(cfg)
        submitted, failed = 0, 0
        errors = []
        for inc in submittable:
            try:
                body = {
                    "base": {
                        "department_neris_id": dept_neris_id,
                        "incident_number": str(inc.get("neris_id_incident") or inc.get("incident_id") or ""),
                        "location": {"address": inc.get("address", "")},
                    },
                    "incident_types": [{"type": inc.get("incident_type", "OTHER")}],
                    "dispatch": {
                        "incident_number": str(inc.get("neris_id_incident") or inc.get("incident_id") or ""),
                        "call_arrival": inc.get("call_create", ""),
                        "call_answered": inc.get("call_create", ""),
                    },
                }
                client.create_incident(dept_neris_id, body)
                submitted += 1
            except Exception as ex:
                failed += 1
                errors.append(str(ex)[:120])
        return JSONResponse({
            "ok": True, "submitted": submitted, "failed": failed,
            "blocked": pre["with_errors"], "errors": errors[:5],
            "dept_name": dept_name,
        })
    except Exception as e:
        return JSONResponse({"ok": False, "message": str(e)}, status_code=500)


# ── NERIS API Connection ──────────────────────────────────────────────────────

@app.post("/neris/connect", response_class=JSONResponse)
async def neris_connect(
    client_id:     str = Form(...),
    client_secret: str = Form(...),
):
    global _neris_creds
    try:
        import httpx
        token_url = "https://auth.neris.fsri.org/oauth2/token"
        resp = httpx.post(token_url, data={
            "grant_type":    "client_credentials",
            "client_id":     client_id,
            "client_secret": client_secret,
        }, timeout=10)
        if resp.status_code == 200:
            _neris_creds = {
                "client_id":     client_id,
                "client_secret": client_secret,
                "connected":     True,
                "entity":        client_id,
            }
            return JSONResponse({"ok": True, "message": "Connected to NERIS API successfully."})
        else:
            _neris_creds["connected"] = False
            return JSONResponse({"ok": False, "message": f"Authentication failed (HTTP {resp.status_code}). Check your Client ID and Secret."}, status_code=400)
    except Exception as e:
        _neris_creds["connected"] = False
        return JSONResponse({"ok": False, "message": str(e)}, status_code=500)


@app.get("/neris/status", response_class=JSONResponse)
async def neris_status():
    return JSONResponse({
        "connected": _neris_creds.get("connected", False),
        "entity":    _neris_creds.get("entity"),
    })


# ── Resources / Education ─────────────────────────────────────────────────────

@app.get("/resources", response_class=HTMLResponse)
async def resources_index(request: Request):
    from articles import ARTICLES
    return templates.TemplateResponse("resources.html", {
        "request": request,
        "articles": ARTICLES,
    })


@app.get("/neris", response_class=HTMLResponse)
async def neris_education(request: Request):
    from articles import ARTICLES
    neris_articles = [a for a in ARTICLES if a["category"] == "NERIS"]
    return templates.TemplateResponse("neris_education.html", {
        "request": request,
        "neris_articles": neris_articles,
    })


@app.get("/pricing", response_class=HTMLResponse)
async def pricing_page(request: Request):
    return templates.TemplateResponse("pricing.html", {"request": request})


@app.get("/report", response_class=HTMLResponse)
async def report_editor(request: Request, d: str = ""):
    import base64 as _b64, json as _json2, zlib as _zlib
    try:
        _raw = _b64.urlsafe_b64decode(d.encode() + b"==")  # pad tolerance
        data = _json2.loads(_zlib.decompress(_raw).decode())
    except Exception:
        data = {"type": "report", "dept": "Your Department", "chief": "Chief", "subject": "Report", "content": ""}
    return templates.TemplateResponse("report_editor.html", {"request": request, "data": data})


@app.get("/resources/{slug}", response_class=HTMLResponse)
async def resource_article(request: Request, slug: str):
    from articles import get_article, ARTICLES
    article = get_article(slug)
    if not article:
        return HTMLResponse("<h1>Article not found</h1>", status_code=404)
    slugs = [a["slug"] for a in ARTICLES]
    idx = slugs.index(slug)
    prev_article = ARTICLES[idx - 1] if idx > 0 else None
    next_article = ARTICLES[idx + 1] if idx < len(ARTICLES) - 1 else None
    return templates.TemplateResponse("article.html", {
        "request": request,
        "article": article,
        "prev_article": prev_article,
        "next_article": next_article,
    })


# ── Directory ─────────────────────────────────────────────────────────────────

import json as _json
from pathlib import Path as _Path

_DEPT_DIR = _Path(__file__).parent / "data" / "departments"

# ── Department index (one file, loaded once at startup) ───────────────────────
_INDEX_PATH = _Path(__file__).parent / "data" / "departments_index.json"
try:
    _DEPT_INDEX: list[dict] = _json.loads(_INDEX_PATH.read_text())
except Exception:
    _DEPT_INDEX = []

# Lazy per-request cache for full dept data (individual files loaded on demand)
_DEPT_CACHE: dict[str, dict] = {}

def _load_dept(path: _Path) -> dict:
    stem = path.stem
    if stem in _DEPT_CACHE:
        return _DEPT_CACHE[stem]
    try:
        d = _json.loads(path.read_text())
        _DEPT_CACHE[stem] = d
        return d
    except Exception:
        return {}

STATE_NAMES = {
    "AL": "Alabama", "AK": "Alaska", "AZ": "Arizona", "AR": "Arkansas",
    "CA": "California", "CO": "Colorado", "CT": "Connecticut", "DE": "Delaware",
    "FL": "Florida", "GA": "Georgia", "HI": "Hawaii", "ID": "Idaho",
    "IL": "Illinois", "IN": "Indiana", "IA": "Iowa", "KS": "Kansas",
    "KY": "Kentucky", "LA": "Louisiana", "ME": "Maine", "MD": "Maryland",
    "MA": "Massachusetts", "MI": "Michigan", "MN": "Minnesota", "MS": "Mississippi",
    "MO": "Missouri", "MT": "Montana", "NE": "Nebraska", "NV": "Nevada",
    "NH": "New Hampshire", "NJ": "New Jersey", "NM": "New Mexico", "NY": "New York",
    "NC": "North Carolina", "ND": "North Dakota", "OH": "Ohio", "OK": "Oklahoma",
    "OR": "Oregon", "PA": "Pennsylvania", "RI": "Rhode Island", "SC": "South Carolina",
    "SD": "South Dakota", "TN": "Tennessee", "TX": "Texas", "UT": "Utah",
    "VT": "Vermont", "VA": "Virginia", "WA": "Washington", "WV": "West Virginia",
    "WI": "Wisconsin", "WY": "Wyoming", "DC": "Washington D.C.",
}

INCIDENT_TYPE_LABELS = {
    "structure_fire": "Structure Fire",
    "vehicle_fire": "Vehicle Fire",
    "wildland_fire": "Wildland Fire",
    "hazmat": "Hazmat",
    "rescue_ems": "Rescue / EMS",
    "false_alarm": "False Alarm",
    "other_fire": "Other Fire",
    "other": "Other",
}


@app.get("/api/dept-search")
async def dept_search(q: str = "", limit: int = 8):
    q_lower = q.strip().lower()
    if len(q_lower) < 2:
        return []
    results = []
    for entry in _DEPT_INDEX:
        if q_lower in entry["name"].lower() or q_lower in entry["city"].lower():
            results.append({
                "name":  entry["name"],
                "city":  entry["city"],
                "state": entry["state"].lower(),
                "fdid":  entry["fdid"],
            })
    results.sort(key=lambda x: (not x["name"].lower().startswith(q_lower), x["name"].lower()))
    return results[:limit]


@app.get("/directory", response_class=HTMLResponse)
async def directory_index(request: Request):
    states = {}
    for entry in _DEPT_INDEX:
        code = entry.get("state", "").upper()
        if not code:
            continue
        if code not in states:
            states[code] = {"code": code, "name": STATE_NAMES.get(code, code), "dept_count": 0, "incident_count": 0}
        states[code]["dept_count"] += 1
        states[code]["incident_count"] += entry.get("total_incidents", 0)

    state_list = sorted(states.values(), key=lambda s: s["name"])
    total_depts = sum(s["dept_count"] for s in state_list)
    total_incidents = sum(s["incident_count"] for s in state_list)

    response = templates.TemplateResponse("directory_index.html", {
        "request": request,
        "states": state_list,
        "total_depts": f"{total_depts:,}",
        "total_incidents": f"{total_incidents:,}",
        "total_states": len(state_list),
    })
    response.headers["Cache-Control"] = "public, max-age=3600, s-maxage=86400"
    return response


@app.get("/directory/{state_code}", response_class=HTMLResponse)
async def directory_state(request: Request, state_code: str):
    state_code = state_code.upper()
    state_name = STATE_NAMES.get(state_code, state_code)

    prefix = f"{state_code.lower()}-"
    departments = []
    max_incidents = 1
    for entry in _DEPT_INDEX:
        if entry.get("stem", "").startswith(prefix) or entry.get("state", "").upper() == state_code:
            inc = entry.get("total_incidents", 0)
            max_incidents = max(max_incidents, inc)
            top_type = ""
            if entry.get("incident_types"):
                top_key = max(entry["incident_types"], key=lambda k: entry["incident_types"][k])
                top_type = INCIDENT_TYPE_LABELS.get(top_key, top_key)
            departments.append({**entry, "top_type": top_type})

    departments.sort(key=lambda d: d.get("total_incidents", 0), reverse=True)

    if not departments:
        return HTMLResponse(f"<h1>No data found for {state_name}</h1>", status_code=404)

    total_incidents = sum(d.get("total_incidents", 0) for d in departments)
    with_data = sum(1 for d in departments if d.get("total_incidents", 0) > 0)

    response = templates.TemplateResponse("directory_state.html", {
        "request": request,
        "state_code": state_code.lower(),
        "state_name": state_name,
        "departments": departments,
        "total": len(departments),
        "total_incidents": total_incidents,
        "with_data": with_data,
        "max_incidents": max_incidents,
    })
    response.headers["Cache-Control"] = "public, max-age=3600, s-maxage=86400"
    return response


@app.get("/directory/{state_code}/{fdid}", response_class=HTMLResponse)
async def directory_dept(request: Request, state_code: str, fdid: str):
    state_code = state_code.upper()
    stem = f"{state_code.lower()}-{fdid}"
    dept = _DEPT_CACHE.get(stem) or _load_dept(_DEPT_DIR / f"{stem}.json")
    if not dept:
        return HTMLResponse("<h1>Department not found</h1>", status_code=404)

    # Monthly chart data
    monthly = dept.get("monthly_volume", {})
    monthly_labels = list(monthly.keys())
    monthly_values = list(monthly.values())

    # Incident type breakdown
    inc_types = dept.get("incident_types", {})
    total = sum(inc_types.values()) or 1
    incident_types = sorted([
        {
            "key": k,
            "label": INCIDENT_TYPE_LABELS.get(k, k),
            "count": v,
            "pct": round(v * 100 / total),
        }
        for k, v in inc_types.items()
    ], key=lambda x: x["count"], reverse=True)

    # Fire vs. non-fire split
    fire_keys = {"structure_fire", "vehicle_fire", "wildland_fire", "other_fire"}
    fire_count = sum(v for k, v in inc_types.items() if k in fire_keys)
    non_fire_count = total - fire_count

    # Incidents per firefighter
    career = int(dept.get("num_ff_career") or 0)
    volunteer = int(dept.get("num_ff_volunteer") or 0)
    total_ff = career + volunteer
    incidents_per_ff = round(dept.get("total_incidents", 0) / total_ff, 1) if total_ff else None

    # Incidents per station
    num_stations = int(dept.get("num_stations") or 0)
    incidents_per_station = round(dept.get("total_incidents", 0) / num_stations, 1) if num_stations else None

    # Incident rate per 1,000 residents (combines incident + census)
    census = dept.get("census") or {}
    population = census.get("population") or 0
    incident_rate_per_1k = round(dept.get("total_incidents", 0) / population * 1000, 1) if population else None

    # Seasonal totals (Q1–Q4) from monthly volume
    quarters = {"Q1": 0, "Q2": 0, "Q3": 0, "Q4": 0}
    month_to_q = {"01":"Q1","02":"Q1","03":"Q1","04":"Q2","05":"Q2","06":"Q2",
                  "07":"Q3","08":"Q3","09":"Q3","10":"Q4","11":"Q4","12":"Q4"}
    for key, val in monthly.items():
        mm = key[5:7] if len(key) >= 7 else key[-2:]
        q = month_to_q.get(mm)
        if q:
            quarters[q] += val

    # Peak month
    month_names = {"01":"January","02":"February","03":"March","04":"April","05":"May",
                   "06":"June","07":"July","08":"August","09":"September",
                   "10":"October","11":"November","12":"December"}
    peak_month_key = max(monthly, key=monthly.get) if monthly else None
    peak_month_count = monthly.get(peak_month_key) if peak_month_key else None
    peak_month_name = month_names.get(peak_month_key[5:7] if peak_month_key and len(peak_month_key) >= 7 else "", "") if peak_month_key else None

    layout = request.query_params.get("v", "")
    template_map = {"1":"directory_dept_v1.html","2":"directory_dept_v2.html","3":"directory_dept_v3.html","4":"directory_dept_v4.html","5":"directory_dept_v5.html"}
    template_name = template_map.get(layout, "directory_dept.html")
    if dept.get("total_incidents", 0) < 10:
        template_name = "directory_dept_light.html"

    response = templates.TemplateResponse(template_name, {
        "request": request,
        "dept": dept,
        "state_name": STATE_NAMES.get(state_code, state_code),
        "monthly_labels": monthly_labels,
        "monthly_values": monthly_values,
        "incident_types": incident_types,
        "fire_count": fire_count,
        "non_fire_count": non_fire_count,
        "total_ff": total_ff,
        "incidents_per_ff": incidents_per_ff,
        "incidents_per_station": incidents_per_station,
        "incident_rate_per_1k": incident_rate_per_1k,
        "quarters": quarters,
        "peak_month_name": peak_month_name,
        "peak_month_count": peak_month_count,
        "census": dept.get("census"),
        "disasters": dept.get("disasters"),
        "peers": dept.get("peers"),
        "nri": dept.get("nri"),
    })
    response.headers["Cache-Control"] = "public, max-age=3600, s-maxage=86400"
    return response


@app.post("/convert/download")
async def convert_download(
    content: str = Form(...),
    filename: str = Form("neris-incidents.json"),
):
    return PlainTextResponse(
        content,
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
        media_type="application/json",
    )


# ── Generate & Email Report ───────────────────────────────────────────────────

def _markdown_to_html(md: str) -> str:  # also registered as Jinja2 filter below
    """Very lightweight markdown → HTML for email bodies."""
    import re
    lines = md.split("\n")
    out = []
    for line in lines:
        if line.startswith("### "):
            out.append(f"<h3 style='color:#1a1a1a;margin:20px 0 6px'>{line[4:]}</h3>")
        elif line.startswith("## "):
            out.append(f"<h2 style='color:#1a1a1a;margin:24px 0 8px;border-bottom:2px solid #D63737;padding-bottom:6px'>{line[3:]}</h2>")
        elif line.startswith("# "):
            out.append(f"<h1 style='color:#D63737;margin:0 0 16px'>{line[2:]}</h1>")
        elif line.startswith("- ") or line.startswith("* "):
            out.append(f"<li style='margin-bottom:4px'>{line[2:]}</li>")
        elif line.strip() == "":
            out.append("<br/>")
        else:
            line = re.sub(r"\*\*(.+?)\*\*", r"<strong>\1</strong>", line)
            out.append(f"<p style='margin:6px 0;line-height:1.6'>{line}</p>")
    return "\n".join(out)


templates.env.filters["markdown_to_html"] = _markdown_to_html


def _load_sample_incidents():
    """Load incidents from sample_nfirs_large.csv, falling back to mock data."""
    import os
    csv_path = os.path.join(os.path.dirname(__file__), "sample_nfirs_large.csv")
    if os.path.exists(csv_path):
        try:
            with open(csv_path, encoding="utf-8", errors="replace") as f:
                csv_text = f.read()
            incidents, _ = convert_nfirs_csv(csv_text)
            if incidents:
                return incidents
        except Exception:
            pass
    from mock_data import generate_incidents
    return generate_incidents()


def _save_lead(name: str, email: str, dept_name: str, dept_state: str, report_type: str, source: str = "onboarding"):
    """Save lead to Supabase. Non-blocking — errors are swallowed."""
    try:
        from supabase import create_client
        sb = create_client(SUPABASE_URL, SUPABASE_SERVICE_ROLE_KEY)
        sb.table("5alarmdata_leads").insert({
            "name":        name or None,
            "email":       email,
            "dept_name":   dept_name or None,
            "dept_state":  dept_state or None,
            "report_type": report_type or None,
            "source":      source,
        }).execute()
    except Exception:
        pass


@app.post("/generate-email", response_class=JSONResponse)
async def generate_and_email(
    report_type:   str = Form("report"),
    email:         str = Form(...),
    name:          str = Form(""),
    dept_name:     str = Form(""),
    dept_state:    str = Form(""),
    chief_name:    str = Form(""),
    grant_request: str = Form(""),
):
    import httpx
    try:
        incidents = _load_sample_incidents()
        dept  = dept_name.strip() or "Your Fire Department"
        chief = chief_name.strip() or name.strip() or "Chief"
        stats = summarize_incidents(incidents)
        period_label = "the past year"

        if report_type in ("board_report", "report", ""):
            result  = generate_chiefs_report(dept, stats, period_label)
            subject = f"Your 5AlarmData Board Report — {dept}"
        elif report_type == "trends":
            result  = generate_trend_summary(dept, stats, period_label)
            subject = f"Your 5AlarmData Incident Trend Summary — {dept}"
        elif report_type == "grant":
            result  = generate_grant_narrative(
                dept_name=dept, stats=stats, grant_type="AFG",
                request_description=grant_request or "firefighting equipment and personnel",
                period=period_label,
            )
            subject = f"Your 5AlarmData Grant Narrative — {dept}"
        elif report_type == "iso":
            from insights.iso_narrative import compute_iso_metrics, generate_iso_narrative
            iso_metrics = compute_iso_metrics(incidents)
            result  = generate_iso_narrative(dept, stats, iso_metrics, period_label)
            subject = f"Your 5AlarmData ISO Evidence Pack — {dept}"
        elif report_type == "staffing":
            from staffing.analyzer import analyze_staffing
            from insights.staffing_narrative import generate_staffing_narrative
            rpt    = analyze_staffing(incidents)
            result = generate_staffing_narrative(dept, rpt)
            subject = f"Your 5AlarmData Staffing Justification — {dept}"
        elif report_type == "compliance":
            from compliance.checker import check_compliance
            import anthropic, json as _json
            rpt = check_compliance(incidents)
            _client = anthropic.Anthropic(api_key=ANTHROPIC_API_KEY)
            _msg = _client.messages.create(
                model=ANTHROPIC_MODEL,
                max_tokens=1200,
                messages=[{"role": "user", "content": f"""You are a NERIS compliance specialist reviewing a fire department's incident data against the mandatory NERIS reporting requirements (deadline: January 2026, already passed).

Compliance Report for {dept}:
{_json.dumps(rpt, indent=2)}

Write a plain-English NERIS compliance report the fire chief can act on immediately. Structure it as:

## NERIS Compliance Summary
One paragraph — overall compliance percentage, what it means practically, and urgency given the deadline has passed.

## Module-by-Module Results
For each module, state the compliance percentage, whether it passes, and what specific fields are missing.

## Top Gaps to Fix
For the worst gaps, explain in plain English what the missing field is, why NERIS requires it, and what the chief should tell their RMS vendor or records staff to do.

## Next Steps
Three specific actions the department should take this week.

Use plain language. No jargon. Be direct about any critical gaps."""}],
            )
            result  = _msg.content[0].text
            subject = f"Your 5AlarmData NERIS Compliance Check — {dept}"
        else:
            result  = generate_chiefs_report(dept, stats, period_label)
            subject = f"Your 5AlarmData Report — {dept}"

        html_body = f"""
<div style="font-family:Georgia,serif;max-width:680px;margin:0 auto;padding:32px 24px;background:#ffffff;color:#1a1a1a">
  <div style="background:#D63737;padding:20px 28px;border-radius:8px 8px 0 0">
    <div style="font-size:22px;font-weight:700;color:#fff;font-family:Georgia,serif">5AlarmData</div>
    <div style="font-size:13px;color:rgba(255,255,255,.8);margin-top:4px">Fire Department Report</div>
  </div>
  <div style="background:#f8f8f8;padding:28px 32px;border:1px solid #ddd;border-top:none;border-radius:0 0 8px 8px">
    <p style="font-size:16px;color:#333;line-height:1.6;margin:0 0 28px">
      Hi {chief},<br/><br/>
      Here is your <strong>{report_type.replace("_"," ").title()}</strong> for <strong>{dept}</strong>.
      This sample was built from a real fire department dataset so you can see exactly what the finished product looks like.
      To get a report built from your department's own numbers, create a free account and connect your data.
    </p>
    <p style="margin:0 0 28px">
      <a href="https://www.5alarmdata.com/sign-in" style="display:inline-block;padding:12px 24px;background:#D63737;color:#fff;border-radius:6px;font-size:15px;font-weight:700;text-decoration:none;font-family:Georgia,serif">Create Your Free Account →</a>
    </p>
    <div style="background:#fff;border:1px solid #ddd;border-radius:6px;padding:28px 32px;line-height:1.7">
      {_markdown_to_html(result)}
    </div>
    <p style="font-size:13px;color:#999;margin:24px 0 0;border-top:1px solid #eee;padding-top:16px">
      Questions? Reply to this email or write to
      <a href="mailto:info@5alarmdata.com" style="color:#D63737">info@5alarmdata.com</a>.
    </p>
  </div>
</div>"""

        # Build report URL (compressed + base64-encoded content, no DB needed)
        import base64 as _b64, json as _json2, zlib as _zlib
        _report_data = _json2.dumps({
            "type":    report_type,
            "dept":    dept,
            "chief":   chief,
            "subject": subject,
            "content": result,
        })
        _compressed = _zlib.compress(_report_data.encode(), level=9)
        _encoded = _b64.urlsafe_b64encode(_compressed).decode()
        report_url = f"/report?d={_encoded}"

        # Save lead to Supabase
        _save_lead(name, email, dept_name, dept_state, report_type, "onboarding")

        # Send report to user — use verified fireinsight.app sender domain
        httpx.post(
            "https://api.resend.com/emails",
            headers={"Authorization": f"Bearer {RESEND_API_KEY}", "Content-Type": "application/json"},
            json={
                "from":    "5AlarmData <info@solofi.io>",
                "to":      [email],
                "subject": subject,
                "html":    html_body,
            },
            timeout=20,
        )

        # Notify admin
        httpx.post(
            "https://api.resend.com/emails",
            headers={"Authorization": f"Bearer {RESEND_API_KEY}", "Content-Type": "application/json"},
            json={
                "from":    "5AlarmData <info@solofi.io>",
                "to":      [DEMO_EMAIL_TO],
                "subject": f"[5AlarmData Lead] {name or email} — {dept} ({report_type})",
                "text":    f"Name: {name}\nEmail: {email}\nDept: {dept}\nState: {dept_state}\nReport: {report_type}",
            },
            timeout=10,
        )

        return JSONResponse({"ok": True, "report_url": report_url})

    except Exception as e:
        return JSONResponse({"ok": False, "error": str(e)}, status_code=500)


# ── Contact Form ──────────────────────────────────────────────────────────────

@app.post("/contact", response_class=JSONResponse)
async def contact_form(
    name:    str = Form(""),
    email:   str = Form(...),
    message: str = Form(""),
):
    try:
        import httpx
        body = f"Contact form submission from 5AlarmData landing page:\n\nName:    {name or 'Not provided'}\nEmail:   {email}\nMessage: {message or 'No message'}"
        httpx.post(
            "https://api.resend.com/emails",
            headers={"Authorization": f"Bearer {RESEND_API_KEY}", "Content-Type": "application/json"},
            json={
                "from":    "5AlarmData <info@solofi.io>",
                "to":      [DEMO_EMAIL_TO],
                "subject": f"[5AlarmData Contact] {name or email}",
                "text":    body,
            },
            timeout=10,
        )
        _save_lead(name, email, "", "", "contact", "contact_form")
        return JSONResponse({"ok": True})
    except Exception as e:
        return JSONResponse({"ok": False, "error": str(e)}, status_code=500)


# ── Demo Request ──────────────────────────────────────────────────────────────

@app.post("/demo", response_class=JSONResponse)
async def demo_request(
    name:        str = Form(...),
    email:       str = Form(...),
    department:  str = Form(...),
    role:        str = Form(""),
    report_type: str = Form(""),
):
    try:
        import httpx
        body = f"""New sample report request from 5AlarmData:\n
Name:         {name}
Email:        {email}
Department:   {department}
Role:         {role or "Not specified"}
Report Type:  {report_type or "Not specified"}
"""
        resp = httpx.post(
            "https://api.resend.com/emails",
            headers={"Authorization": f"Bearer {RESEND_API_KEY}", "Content-Type": "application/json"},
            json={
                "from":    "5AlarmData <info@solofi.io>",
                "to":      [DEMO_EMAIL_TO],
                "subject": f"Sample Report Request: {department}",
                "text":    body,
            },
            timeout=10,
        )
        # Also save lead to Supabase
        _save_lead(name, email, department, "", report_type, "landing_modal")
        if resp.status_code in (200, 201):
            return JSONResponse({"ok": True})
        else:
            return JSONResponse({"ok": False, "error": f"Email send failed ({resp.status_code})"}, status_code=500)
    except Exception as e:
        return JSONResponse({"ok": False, "error": str(e)}, status_code=500)


# ── Custom Report Builder ─────────────────────────────────────────────────────

@app.post("/report-builder", response_class=HTMLResponse)
async def report_builder(
    request:   Request,
    neris_id:  str  = Form("MOCK-001"),
    use_mock:  bool = Form(False),
    start:     str  = Form(""),
    end:       str  = Form(""),
    compare_start: str = Form(""),
    compare_end:   str = Form(""),
    metrics:   str  = Form(""),   # comma-separated list
    period_label: str = Form(""),
):
    error  = None
    report = None
    dept_name = ""

    try:
        from report_builder.builder import build_custom_report
        incidents, dept_name = load_incidents(neris_id, use_mock, start, end)

        compare_incidents = []
        if compare_start and compare_end:
            compare_incidents, _ = load_incidents(neris_id, use_mock, compare_start, compare_end)

        selected_metrics = [m.strip() for m in metrics.split(",") if m.strip()] if metrics else []

        if not incidents:
            error = "No incidents found for the given parameters."
        else:
            report = build_custom_report(
                dept_name=dept_name,
                incidents=incidents,
                compare_incidents=compare_incidents,
                metrics=selected_metrics,
                period_label=period_label or (f"{start} to {end}" if start and end else "the past year"),
            )
    except Exception as e:
        error = str(e)

    return templates.TemplateResponse("index.html", {
        "request":    request,
        "active_tab": "report_builder",
        "report_builder": {
            "dept_name": dept_name,
            "report":    report,
            "error":     error,
            "form": {
                "neris_id":     neris_id,
                "start":        start,
                "end":          end,
                "compare_start": compare_start,
                "compare_end":   compare_end,
                "metrics":      metrics,
                "period_label": period_label,
            },
        },
    })


@app.post("/report-builder/export")
async def report_builder_export(
    content:  str = Form(...),
    filename: str = Form("custom-report.md"),
):
    return PlainTextResponse(
        content,
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
        media_type="text/markdown",
    )


# ── State NERIS Coordinator Export ───────────────────────────────────────────

@app.post("/state-export", response_class=HTMLResponse)
async def state_export(
    request:   Request,
    neris_id:  str  = Form("MOCK-001"),
    use_mock:  bool = Form(False),
    start:     str  = Form(""),
    end:       str  = Form(""),
    state:     str  = Form("VA"),
):
    error     = None
    result    = None
    dept_name = ""

    try:
        from exports.state_coordinator import generate_state_export
        incidents, dept_name = load_incidents(neris_id, use_mock, start, end)
        if not incidents:
            error = "No incidents found for the given parameters."
        else:
            result = generate_state_export(incidents, state, dept_name)
    except Exception as e:
        error = str(e)

    return templates.TemplateResponse("index.html", {
        "request":    request,
        "active_tab": "state_export",
        "state_export": {
            "dept_name": dept_name,
            "result":    result,
            "error":     error,
            "form": {"state": state, "start": start, "end": end},
        },
    })


@app.post("/state-export/download")
async def state_export_download(
    content:  str = Form(...),
    filename: str = Form("state-neris-export.csv"),
):
    return PlainTextResponse(
        content,
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
        media_type="text/csv",
    )


# ── Historical NFIRS Archive Analyzer ────────────────────────────────────────

@app.post("/archive", response_class=HTMLResponse)
async def archive_analysis(
    request: Request,
    file:    UploadFile = File(...),
    neris_id: str  = Form("MOCK-001"),
    use_mock: bool = Form(False),
):
    error     = None
    result    = None
    dept_name = ""
    filename  = file.filename or "upload.csv"

    try:
        from archive.nfirs_analyzer import analyze_nfirs_archive
        from mock_data import DEPT
        raw_bytes = await file.read()
        csv_text  = None
        for enc in ("utf-8", "latin-1", "cp1252"):
            try:
                csv_text = raw_bytes.decode(enc)
                break
            except UnicodeDecodeError:
                continue
        if csv_text is None:
            raise ValueError("Could not decode file. Please export as UTF-8 or Latin-1 CSV.")

        # Load current NERIS data for comparison baseline
        current_incidents, dept_name = resolve_incidents(request, "MOCK-001", use_mock, "", "")

        result = analyze_nfirs_archive(csv_text, current_incidents, dept_name)

    except Exception as e:
        error = str(e)

    return templates.TemplateResponse("index.html", {
        "request":    request,
        "active_tab": "archive",
        "archive": {
            "dept_name": dept_name,
            "filename":  filename,
            "result":    result,
            "error":     error,
        },
    })


@app.get("/sitemap.xml")
async def sitemap():
    from articles import ARTICLES
    base = "https://www.5alarmdata.com"
    static_urls = [
        ("", "weekly", "1.0"),
        ("/about", "monthly", "0.8"),
        ("/resources", "weekly", "0.9"),
        ("/neris", "monthly", "0.8"),
        ("/directory", "monthly", "0.7"),
        ("/start", "monthly", "0.6"),
    ]
    urls = []
    for path, changefreq, priority in static_urls:
        urls.append(
            f"  <url>\n"
            f"    <loc>{base}{path}</loc>\n"
            f"    <changefreq>{changefreq}</changefreq>\n"
            f"    <priority>{priority}</priority>\n"
            f"  </url>"
        )
    for article in ARTICLES:
        urls.append(
            f"  <url>\n"
            f"    <loc>{base}/resources/{article['slug']}</loc>\n"
            f"    <changefreq>monthly</changefreq>\n"
            f"    <priority>0.7</priority>\n"
            f"  </url>"
        )
    xml = (
        '<?xml version="1.0" encoding="UTF-8"?>\n'
        '<urlset xmlns="http://www.sitemaps.org/schemas/sitemap/0.9">\n'
        + "\n".join(urls)
        + "\n</urlset>"
    )
    return Response(content=xml, media_type="application/xml")
